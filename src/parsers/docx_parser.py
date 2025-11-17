"""DOCX document parser."""

import logging
from pathlib import Path
from typing import List

from docx import Document as DocxDocument
from docx.table import Table

from .base import BaseParser, Document

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """Parser for DOCX documents."""

    def parse(self, file_path: Path) -> List[Document]:
        """
        Parse a DOCX document.

        Args:
            file_path: Path to the DOCX file

        Returns:
            List containing a single Document object
        """
        try:
            doc = DocxDocument(str(file_path))

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            text = "\n".join(paragraphs)
            text = self._clean_text(text)

            # Extract metadata
            core_properties = doc.core_properties
            metadata = {
                "file_name": file_path.name,
                "file_path": str(file_path),
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "subject": core_properties.subject or "",
                "created": str(core_properties.created) if core_properties.created else "",
                "modified": str(core_properties.modified) if core_properties.modified else "",
            }

            # Extract tables if enabled
            tables = []
            if self.extract_tables and doc.tables:
                tables = self._extract_tables(doc.tables)

            # Create document
            document = Document(
                text=text,
                metadata=metadata,
                tables=tables,
                source=str(file_path),
            )

            logger.info(f"Parsed DOCX document: {file_path.name}")
            return [document]

        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise

    def _extract_tables(self, tables: List[Table]) -> List[dict]:
        """
        Extract tables from DOCX document.

        Args:
            tables: List of table objects

        Returns:
            List of table dictionaries
        """
        extracted_tables = []

        for table_idx, table in enumerate(tables):
            table_data = []
            headers = []

            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]

                # First row is typically headers
                if row_idx == 0:
                    headers = row_data
                else:
                    table_data.append(row_data)

            if table_data:
                extracted_tables.append(
                    {
                        "table_id": table_idx,
                        "headers": headers,
                        "data": table_data,
                        "num_rows": len(table_data),
                        "num_cols": len(headers) if headers else len(table_data[0]) if table_data else 0,
                    }
                )

        logger.info(f"Extracted {len(extracted_tables)} tables")
        return extracted_tables
